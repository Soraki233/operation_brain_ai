from __future__ import annotations

import concurrent.futures
import logging
import re
import zipfile
from pathlib import Path
from typing import List
from xml.etree import ElementTree as ET

import pandas as pd

logger = logging.getLogger(__name__)
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import docx as _docx
import docx2txt
from langchain_community.document_loaders import PyPDFLoader

from service.vector_store_service import VectorStoreService
from service.structure_analyzer import StructureAnalyzer, StructureAnalyzeError
from core.settings import settings


class KnowledgeIngestService:
    def __init__(self) -> None:
        self.vector_store = VectorStoreService().get_store()
        # 懒加载：仅在真正需要结构化切分时才初始化 LLM 连接
        self._structure_analyzer: StructureAnalyzer | None = None

    def _get_structure_analyzer(self) -> StructureAnalyzer:
        if self._structure_analyzer is None:
            self._structure_analyzer = StructureAnalyzer()
        return self._structure_analyzer

    def load_file(self, file_path: str) -> List[Document]:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            docs = PyPDFLoader(str(path)).load()
        elif suffix in {".docx", ".doc"}:
            docs = self._load_word_doc(path)
        elif suffix in {".xlsx", ".xls"}:
            # 避开 UnstructuredExcelLoader（其 xlsx 分支依赖 networkx / unstructured 全家桶），
            # 直接用 pandas 按 sheet 读取。不用 CSV：对 LLM 来说 ",,,," 这种空列
            # 几乎没有语义，改成 "字段=值 | 字段=值" 的记录格式，召回片段能直接
            # 被模型理解，表头也会被切片逻辑自动复用到每个 chunk。
            engine = "openpyxl" if suffix == ".xlsx" else "xlrd"
            sheet_names = pd.ExcelFile(str(path), engine=engine).sheet_names
            docs = []
            for sheet_name in sheet_names:
                df = self._read_excel_sheet_smart(
                    str(path), sheet_name=sheet_name, engine=engine
                )
                df = self._clean_excel_dataframe(df)
                if df.empty:
                    continue
                row_docs = self._dataframe_to_row_documents(
                    df, sheet_name=sheet_name, source=str(path)
                )
                docs.extend(row_docs)
        elif suffix in {".md", ".txt"}:
            text = path.read_text(encoding="utf-8", errors="ignore")
            docs = [Document(page_content=text, metadata={"source": str(path)})]
        else:
            raise ValueError(f"暂不支持的文件类型: {suffix}")

        cleaned_docs: List[Document] = []
        for doc in docs:
            content = (doc.page_content or "").strip()
            if not content:
                continue
            cleaned_docs.append(
                Document(
                    page_content=content,
                    metadata=doc.metadata or {},
                )
            )
        return cleaned_docs

    def split_documents(
        self,
        docs: List[Document],
        *,
        kb_id: str,
        file_id: str,
    ) -> List[Document]:
        splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.KNOWLEDGE_CHUNK_SIZE,
            chunk_overlap=settings.KNOWLEDGE_CHUNK_OVERLAP,
            separators=[
                "\n\n",
                "\n",
                "。",
                "！",
                "？",
                "；",
                "，",
                " ",
                "",
                ",",
                "\t",
                "\r",
                " ",
            ],
        )

        chunks: List[Document] = []
        non_excel_docs: List[Document] = []
        for doc in docs:
            metadata = dict(doc.metadata or {})
            metadata.update({"kb_id": kb_id, "file_id": file_id})
            normalized = Document(page_content=doc.page_content, metadata=metadata)
            # row_level=True 的 Document 是由 _dataframe_to_row_documents
            # 生成的行级 chunk，已经是最小单元，直接追加，不再二次切分。
            if metadata.get("row_level"):
                chunks.append(normalized)
            else:
                non_excel_docs.append(normalized)

        if non_excel_docs:
            chunks.extend(
                self._split_non_excel_documents(non_excel_docs, splitter=splitter)
            )

        for idx, chunk in enumerate(chunks):
            chunk.metadata["chunk_index"] = idx

        return chunks

    # 结构化单元超过 chunk_size 的多少倍时触发二次切分
    _STRUCTURE_OVERSIZE_FACTOR: float = 1.5
    # 单文档整体结构分析的硬超时（秒）。超过直接 fallback 到 Recursive 切分，
    # 避免 LLM 异常（挂起 / 反复重试）把后台 ingest 线程卡住导致文件永远
    # 停在 "processing" 状态。
    _STRUCTURE_ANALYZE_HARD_TIMEOUT: int = 300

    def _split_non_excel_documents(
        self,
        docs: List[Document],
        *,
        splitter: RecursiveCharacterTextSplitter,
    ) -> List[Document]:
        """PDF/Word/TXT/MD 走 LLM 结构感知切分；失败或关闭时退化为
        RecursiveCharacterTextSplitter。
        """
        if not docs:
            return []

        if not settings.STRUCTURE_ANALYZER_ENABLED:
            return splitter.split_documents(docs)

        try:
            analyzer = self._get_structure_analyzer()
        except Exception as e:  # noqa: BLE001
            logger.warning(
                "[STRUCTURE] 初始化失败，退化为 Recursive 切分 err=%s", e
            )
            return splitter.split_documents(docs)

        oversize_limit = int(
            settings.KNOWLEDGE_CHUNK_SIZE * self._STRUCTURE_OVERSIZE_FACTOR
        )
        result: List[Document] = []

        for i, doc in enumerate(docs, start=1):
            text = (doc.page_content or "").strip()
            if not text:
                continue
            logger.info(
                "[STRUCTURE] 开始分析文档 %d/%d | source=%s 文本=%dchar hard_timeout=%ds",
                i,
                len(docs),
                doc.metadata.get("source"),
                len(text),
                self._STRUCTURE_ANALYZE_HARD_TIMEOUT,
            )
            try:
                # 整篇文档分析放到独立线程，套一个硬超时，
                # 防止 LLM 异常挂起把 ingest 线程永远阻塞。
                with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
                    future = pool.submit(analyzer.analyze, text)
                    units = future.result(
                        timeout=self._STRUCTURE_ANALYZE_HARD_TIMEOUT
                    )
            except concurrent.futures.TimeoutError:
                logger.warning(
                    "[STRUCTURE] 分析超时(%ds)，回退 Recursive 切分 source=%s",
                    self._STRUCTURE_ANALYZE_HARD_TIMEOUT,
                    doc.metadata.get("source"),
                )
                result.extend(splitter.split_documents([doc]))
                continue
            except StructureAnalyzeError as e:
                logger.warning(
                    "[STRUCTURE] 分析失败，回退 Recursive 切分 source=%s err=%s",
                    doc.metadata.get("source"),
                    e,
                )
                result.extend(splitter.split_documents([doc]))
                continue
            except Exception as e:  # noqa: BLE001
                logger.exception(
                    "[STRUCTURE] 未预期异常，回退 Recursive 切分 err=%s", e
                )
                result.extend(splitter.split_documents([doc]))
                continue
            logger.info(
                "[STRUCTURE] 文档 %d/%d 分析完成 | 得到结构单元=%d",
                i,
                len(docs),
                len(units),
            )

            for unit in units:
                title = unit.title or ""
                u_type = unit.type or "paragraph"
                prefix = (
                    f"【{u_type} · {title}】" if title else f"【{u_type}】"
                )
                page_content = f"{prefix}\n{unit.content}".strip()

                base_meta = dict(doc.metadata or {})
                base_meta.update(
                    {
                        "structure_type": u_type,
                        "structure_title": title,
                        "structure_level": int(unit.level or 0),
                        "keywords": list(unit.keywords or []),
                    }
                )

                # 单个结构单元过大 → 用 Recursive 二次切，metadata 保留结构信息
                if len(page_content) > oversize_limit:
                    sub_docs = splitter.split_documents(
                        [Document(page_content=page_content, metadata=base_meta)]
                    )
                    for sub in sub_docs:
                        merged_meta = dict(base_meta)
                        merged_meta.update(sub.metadata or {})
                        result.append(
                            Document(
                                page_content=sub.page_content,
                                metadata=merged_meta,
                            )
                        )
                else:
                    result.append(
                        Document(page_content=page_content, metadata=base_meta)
                    )

        return result

    # 被视为表头的前缀：遇到以这些前缀开头的行，都会被视为需要复用到每个 chunk 的"表头"。
    _HEADER_LINE_PREFIXES = ("工作表：", "字段：")

    @staticmethod
    def _split_csv_document(doc: Document, *, max_chars: int) -> List[Document]:
        """按行切分表格 Document，保证：
        1. 不会把一行数据切到中间；
        2. 每个 chunk 都重复带上表头行（"工作表：..." / "字段：..."），
           让召回片段是"表头 + 若干完整数据行"，模型能直接读懂字段含义。
        """
        text = doc.page_content or ""
        lines = text.split("\n")
        while lines and not lines[-1].strip():
            lines.pop()

        if not lines:
            return []
        if len(lines) == 1:
            return [doc]

        # 把开头连续的"表头行"识别出来；兼容无前缀的情况（退化成首行当表头）
        header_lines: list[str] = []
        idx = 0
        while idx < len(lines) and lines[idx].startswith(
            KnowledgeIngestService._HEADER_LINE_PREFIXES
        ):
            header_lines.append(lines[idx])
            idx += 1
        if not header_lines:
            header_lines = [lines[0]]
            idx = 1

        header_block = "\n".join(header_lines)
        data_lines = lines[idx:]
        if not data_lines:
            return [doc]

        header_len = len(header_block) + 1
        # 表头本身过长时放弃复用表头，至少不切断行
        carry_header = header_len < max(max_chars // 2, 256)

        chunks: List[Document] = []
        buffer: List[str] = []
        buffer_len = header_len if carry_header else 0
        start_row = 0

        def flush(end_row_exclusive: int) -> None:
            if not buffer:
                return
            body = "\n".join(buffer)
            content = f"{header_block}\n{body}" if carry_header else body
            metadata = dict(doc.metadata or {})
            metadata["row_range"] = [start_row, end_row_exclusive - 1]
            chunks.append(Document(page_content=content, metadata=metadata))

        for i, line in enumerate(data_lines):
            line_len = len(line) + 1
            if buffer and (buffer_len + line_len) > max_chars:
                flush(i)
                buffer = []
                buffer_len = header_len if carry_header else 0
                start_row = i
            buffer.append(line)
            buffer_len += line_len

        flush(len(data_lines))

        return chunks or [doc]

    # 被识别为"时间维度"列的列名关键词（用于 semantic_text 前置时间上下文）
    _DATE_COL_KEYWORDS = ("日期", "时间", "date", "time", "年月", "月份", "年份")

    @staticmethod
    def _dataframe_to_row_documents(
        df: pd.DataFrame,
        *,
        sheet_name: str,
        source: str,
    ) -> List[Document]:
        """把 DataFrame 转成行级 Document 列表。

        每一行数据独立成一个 Document：
        - page_content = semantic_text（语义自然语言，适合向量检索）
        - metadata.row_data = {col: val_str, ...}（原始键值对，不丢字段）

        semantic_text 格式（示例）：
            2024年1月1日 08:00，工作表「电量」：博总线正向有功总功率为 8608.70 kW，
            博总线反向有功总功率为 0 kW（第1行）

        优化点：
        1. 日期列值转成中文年月日（"2026-02-01" → "2026年2月1日"）
        2. 时间维度列（日期/时间）前置作为语义上下文
        3. 列名含括号单位（如 "功率(kW)"）自动拆开写成 "功率为 1500 kW"
        4. 列名中无意义的多级连字符前缀尽量缩短（"博总线-正向有功总" 中保留完整语义）
        """
        if df is None or df.empty:
            return []

        import re as _re
        from datetime import datetime as _datetime

        df = df.copy()

        # ── 日期列检测与格式化 ───────────────────────────────────────────────
        # 两种情况都要处理：
        # 1. pandas 已识别为 datetime64 类型（openpyxl 正常解析时）
        # 2. Excel 存储为数字格式（序列号），pandas 读成 int64/float64，
        #    需要用 origin='1899-12-30' + unit='D' 换算回真实日期。
        #
        # Excel 序列号判定规则：
        #   - 列值全为整数（或整数值的 float）
        #   - 数值范围在 25569~73051（对应 1970-01-01 ~ 2099-12-31）
        #     实际业务数据通常在 40000~55000（2009~2050），放宽到 30000~70000
        #   - 列名包含日期关键词，或整列 ≥90% 落在上述范围内
        _EXCEL_DATE_LO, _EXCEL_DATE_HI = 25000, 70000  # ~1968–2091

        def _looks_like_excel_serial(s: pd.Series) -> bool:
            """判断某列是否像 Excel 日期序列号。"""
            s_clean = s.dropna()
            if len(s_clean) == 0:
                return False
            # 必须是数值
            if not (
                pd.api.types.is_integer_dtype(s_clean)
                or pd.api.types.is_float_dtype(s_clean)
            ):
                return False
            # float 列必须全是整数值（没有小数部分）
            if pd.api.types.is_float_dtype(s_clean):
                if not (s_clean == s_clean.round()).all():
                    return False
            vals = s_clean.astype(float)
            in_range = ((vals >= _EXCEL_DATE_LO) & (vals <= _EXCEL_DATE_HI)).mean()
            return bool(in_range >= 0.9)  # ≥90% 的值在合理日期范围内

        col_is_date_only: dict[str, bool] = {}   # col_name -> True=仅日期 False=带时间
        _date_kws_lower = [kw.lower() for kw in KnowledgeIngestService._DATE_COL_KEYWORDS]

        for col in df.columns:
            s = df[col]
            col_lower = str(col).strip().lower()
            if pd.api.types.is_datetime64_any_dtype(s):
                # 情况 1：pandas 已正确识别 datetime 列
                dt = pd.to_datetime(s, errors="coerce")
                has_time = bool(
                    (
                        (dt.dt.hour.fillna(0) != 0)
                        | (dt.dt.minute.fillna(0) != 0)
                        | (dt.dt.second.fillna(0) != 0)
                    ).any()
                )
                fmt = "%Y-%m-%d %H:%M:%S" if has_time else "%Y-%m-%d"
                df[col] = dt.dt.strftime(fmt)
                col_is_date_only[str(col)] = not has_time
            elif _looks_like_excel_serial(s) and any(
                kw in col_lower for kw in _date_kws_lower
            ):
                # 情况 2：列名含日期关键词 且 列值像 Excel 序列号 → 强制转换
                # Excel 的起始纪元是 1899-12-30（因为 Excel 错误地把 1900-02-29 算进去了）
                try:
                    dt = pd.to_datetime(
                        pd.to_numeric(s, errors="coerce"),
                        unit="D",
                        origin="1899-12-30",
                        errors="coerce",
                    )
                    has_time = bool(
                        (
                            (dt.dt.hour.fillna(0) != 0)
                            | (dt.dt.minute.fillna(0) != 0)
                            | (dt.dt.second.fillna(0) != 0)
                        ).any()
                    )
                    fmt = "%Y-%m-%d %H:%M:%S" if has_time else "%Y-%m-%d"
                    df[col] = dt.dt.strftime(fmt)
                    col_is_date_only[str(col)] = not has_time
                    logger.info(
                        "[EXCEL] 列 '%s' 识别为 Excel 日期序列号，已转换为日期字符串", col
                    )
                except Exception as e:  # noqa: BLE001
                    logger.warning(
                        "[EXCEL] 列 '%s' 序列号转日期失败 err=%s，保留原始值", col, e
                    )

        cols = [str(c).strip() for c in df.columns]

        # ── 列名预处理：拆单位 & 标记时间维度列 ──────────────────────────────
        # col_meta: (raw_col, label, unit, is_time_dim)
        _UNIT_PAT = _re.compile(r"^(.+?)[（(]([^）)]+)[）)]$")
        _date_kws = KnowledgeIngestService._DATE_COL_KEYWORDS

        col_meta: list[tuple[str, str, str, bool]] = []
        for col in cols:
            m = _UNIT_PAT.match(col)
            label = m.group(1).strip() if m else col
            unit  = m.group(2).strip() if m else ""
            is_time = any(kw.lower() in col.lower() for kw in _date_kws)
            col_meta.append((col, label, unit, is_time))

        # ── 值格式化辅助 ─────────────────────────────────────────────────────
        def _fmt_val(val: object) -> str:
            if val is None:
                return ""
            if isinstance(val, float):
                if pd.isna(val):
                    return ""
                if val == int(val):
                    return str(int(val))
                return ("%.6f" % val).rstrip("0").rstrip(".")
            return str(val).strip()

        def _to_chinese_date(val_str: str) -> str:
            """把 "2026-02-01" 或 "2026-02-01 08:00:00" 转成中文年月日。"""
            s = val_str.strip()
            # 尝试日期+时间
            for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
                try:
                    dt = _datetime.strptime(s, fmt)
                    if fmt == "%Y-%m-%d":
                        return f"{dt.year}年{dt.month}月{dt.day}日"
                    # 带时间：如果时间非零保留
                    if dt.hour or dt.minute:
                        return (
                            f"{dt.year}年{dt.month}月{dt.day}日"
                            f" {dt.hour:02d}:{dt.minute:02d}"
                        )
                    return f"{dt.year}年{dt.month}月{dt.day}日"
                except ValueError:
                    continue
            return val_str  # 无法解析则原样返回

        # ── 逐行生成 Document ────────────────────────────────────────────────
        docs: List[Document] = []
        for row_idx, row in enumerate(df.itertuples(index=False, name=None)):
            row_data: dict[str, str] = {}
            time_parts: list[str] = []    # 时间维度值，拼成前缀
            measure_parts: list[str] = [] # 数值/描述字段，拼成主体

            for (raw_col, label, unit, is_time), val in zip(col_meta, row):
                if not raw_col:
                    continue
                val_str = _fmt_val(val)
                if not val_str:
                    continue
                row_data[raw_col] = val_str

                if is_time:
                    # 时间维度列：转成中文日期，直接拼前缀（不写 "XXX为"）
                    cn_val = _to_chinese_date(val_str)
                    time_parts.append(cn_val)
                else:
                    # 数值 / 描述列：写成 "字段为 值 单位"
                    if unit:
                        measure_parts.append(f"{label}为 {val_str} {unit}")
                    else:
                        measure_parts.append(f"{label}为 {val_str}")

            if not row_data:
                continue

            # 组装 semantic_text
            # 格式："{时间上下文}，工作表「X」：{字段描述}（第N行）"
            # 若无时间列，退化为：工作表「X」第N行：{字段描述}
            if time_parts and measure_parts:
                time_ctx = " ".join(time_parts)
                semantic_text = (
                    f"{time_ctx}，工作表「{sheet_name}」："
                    + "，".join(measure_parts)
                    + f"（第{row_idx + 1}行）"
                )
            elif time_parts:
                # 全部字段都是时间维度、没有数值列：只记录时间上下文
                time_ctx = " ".join(time_parts)
                semantic_text = (
                    f"{time_ctx}，工作表「{sheet_name}」（第{row_idx + 1}行）"
                )
            else:
                semantic_text = (
                    f"工作表「{sheet_name}」第{row_idx + 1}行："
                    + "，".join(measure_parts)
                )

            docs.append(
                Document(
                    page_content=semantic_text,
                    metadata={
                        "source": source,
                        "sheet": sheet_name,
                        "row_index": row_idx,
                        "row_level": True,
                        "row_data": row_data,
                    },
                )
            )

        return docs

    # DashScope text-embedding-v3/v4 单次请求最多 10 条文本，超过会直接 4xx。
    # 而 langchain_community.DashScopeEmbeddings.embed_documents 不会主动分批，
    # 因此我们在调用 add_documents 之前手工切片。
    EMBEDDING_BATCH_SIZE: int = 10

    def ingest_file(
        self,
        *,
        kb_id: str,
        file_id: str,
        file_path: str,
    ) -> dict:
        raw_docs = self.load_file(file_path)
        if not raw_docs:
            return {
                "chunk_count": 0,
                "message": "文件没有可入库内容",
            }

        chunks = self.split_documents(
            raw_docs,
            kb_id=kb_id,
            file_id=file_id,
        )

        ids = [f"{file_id}:{i}" for i in range(len(chunks))]

        for start in range(0, len(chunks), self.EMBEDDING_BATCH_SIZE):
            batch_docs = chunks[start : start + self.EMBEDDING_BATCH_SIZE]
            batch_ids = ids[start : start + self.EMBEDDING_BATCH_SIZE]
            try:
                self.vector_store.add_documents(documents=batch_docs, ids=batch_ids)
            except Exception as exc:
                raise self._wrap_embedding_error(exc, batch_index=start) from exc

        return {
            "chunk_count": len(chunks),
            "message": "文件入库成功",
        }

    @staticmethod
    def _wrap_embedding_error(exc: Exception, *, batch_index: int) -> Exception:
        """DashScope 的 HTTPError 会因为 response 是 dict-like 而被
        ``requests.exceptions.HTTPError.__init__`` 二次抛出 ``KeyError('request')``，
        把真正的 status_code/code/message 完全吞掉。这里尽量还原可读信息。
        """
        msg = str(exc) if str(exc) else exc.__class__.__name__

        # 典型表现：KeyError: 'request'  ←  来自 dashscope_response.__getattr__
        if isinstance(exc, KeyError) and "request" in msg:
            msg = (
                "DashScope embedding 请求失败（疑似单批文本数 / 单文本 token 超限，"
                "或 API Key/网络异常）。已自动按 10 条/批切分仍报错，请检查 API "
                "配额或缩小 KNOWLEDGE_CHUNK_SIZE。"
            )

        return RuntimeError(f"[batch_offset={batch_index}] {msg}")

    def delete_file_vectors(self, *, file_id: str, chunk_count: int) -> None:
        if chunk_count <= 0:
            return
        ids = [f"{file_id}:{i}" for i in range(chunk_count)]
        self.vector_store.delete(ids=ids)

    @staticmethod
    def _load_word_doc(path: Path) -> List[Document]:
        """加载 .docx 或 .doc 文件。
        
        尝试顺序：
        1. python-docx（适合 .docx 及以 .doc 保存的 Open XML 文件）
        2. docx2txt（更宽松的解析）
        3. 提示用户转存为 .docx
        """
        text: str | None = None

        # 尝试 1：python-docx（适用标准 Word / 部分 WPS）
        try:
            document = _docx.Document(str(path))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        paragraphs.append(" | ".join(cells))
            text = "\n".join(paragraphs)
        except Exception as e:
            logger.info("[WORD] python-docx 失败，尝试下一种方式 err=%s", e)

        # 尝试 2：docx2txt
        if not text:
            try:
                text = docx2txt.process(str(path))
            except Exception as e:
                logger.info("[WORD] docx2txt 失败，尝试下一种方式 err=%s", e)

        # 尝试 3：直接当 ZIP 解压，从 XML 里提取文本
        # 兼容 WPS 文字、金山 WPS、Office 365、以及各种 relationship 变体
        if not text:
            try:
                text = KnowledgeIngestService._extract_text_from_ooxml_zip(path)
            except Exception as e:
                logger.info("[WORD] ZIP 直解方式失败 err=%s", e)

        if not text or not text.strip():
            suffix = path.suffix.lower()
            logger.error("[WORD] 所有方式均无法提取文本 path=%s", path)
            if suffix == ".doc":
                raise ValueError(
                    "旧版 .doc 文件无法解析，请在 Word 中另存为 .docx 格式后重新上传"
                )
            raise ValueError(
                f"Word 文件解析失败，文件可能为加密或严重损坏: {path.name}"
            )

        return [Document(page_content=text.strip(), metadata={"source": str(path)})]

    @staticmethod
    def _extract_text_from_ooxml_zip(path: Path) -> str:
        """直接把 .docx/.doc 当 ZIP 拆开，从任何包含 w:t 标签的 XML 里提取文本。
        
        可处理：
        - WPS 文字生成的 .docx（relationship namespace 与 Word 不同）
        - 含自定义关系的 Office 365 文件
        - 任意 Open XML 变体
        """
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        text_parts: list[str] = []

        with zipfile.ZipFile(str(path), "r") as zf:
            # 优先读 word/document.xml；其次找所有含 "document" 的 xml
            xml_candidates = [
                n for n in zf.namelist()
                if n.endswith(".xml") and (
                    "document" in n.lower() or
                    "content" in n.lower() or
                    n.startswith("word/")
                )
            ]
            # 如果没有上述匹配，就读所有 xml
            if not xml_candidates:
                xml_candidates = [n for n in zf.namelist() if n.endswith(".xml")]

            for xml_name in xml_candidates:
                try:
                    xml_bytes = zf.read(xml_name)
                    root = ET.fromstring(xml_bytes)
                    # 提取所有 w:t 文本节点
                    for elem in root.iter(f"{{{W_NS}}}t"):
                        if elem.text:
                            text_parts.append(elem.text)
                    # 段落分隔符
                    for elem in root.iter(f"{{{W_NS}}}p"):
                        text_parts.append("\n")
                except Exception:
                    continue

        raw = "".join(text_parts)
        # 合并多余换行
        raw = re.sub(r"\n{3,}", "\n\n", raw).strip()
        return raw

    @staticmethod
    def _read_excel_sheet_smart(
        file_path: str, *, sheet_name: str, engine: str
    ) -> pd.DataFrame:
        """读取单个 sheet，并自动识别多级表头：
        - 规则：扫描前若干行，第一个"多数单元格是数字"的行判为数据行；
          它之前的行都算表头，多级合并成 "主标题-子标题" 这种可读列名。
        - 这样处理图里那种 "日期/时间/博总线(正向有功总/反向有功总)/电量(...)" 的
          合并表头，就不会把 "正向有功总" 这些子标题错误当成数据。
        """
        probe = pd.read_excel(
            file_path, sheet_name=sheet_name, engine=engine, header=None, nrows=12
        )
        header_depth = KnowledgeIngestService._detect_header_depth(probe)

        if header_depth <= 1:
            return pd.read_excel(
                file_path, sheet_name=sheet_name, engine=engine, header=0
            )

        df = pd.read_excel(
            file_path,
            sheet_name=sheet_name,
            engine=engine,
            header=list(range(header_depth)),
        )
        df.columns = [
            KnowledgeIngestService._flatten_header(col) for col in df.columns
        ]
        return df

    @staticmethod
    def _detect_header_depth(probe: pd.DataFrame) -> int:
        """前若干行里，第一个 "数字占比 ≥ 50%" 的行 = 表头深度；否则默认 1。"""
        if probe is None or probe.empty:
            return 1
        for i in range(min(6, len(probe))):
            row = probe.iloc[i]
            non_null = [v for v in row if pd.notna(v)]
            if not non_null:
                continue
            numeric = sum(
                1
                for v in non_null
                if isinstance(v, (int, float)) and not isinstance(v, bool)
            )
            if numeric / len(non_null) >= 0.5:
                return max(1, i)
        return 1

    @staticmethod
    def _flatten_header(col) -> str:
        """把多级表头的 tuple 合并成 "A-B-C"，过滤 NaN / Unnamed / 重复。"""
        if isinstance(col, tuple):
            parts: list[str] = []
            for c in col:
                s = (
                    ""
                    if c is None or (isinstance(c, float) and pd.isna(c))
                    else str(c).strip()
                )
                if not s or s.startswith("Unnamed:"):
                    continue
                if s in parts:
                    continue
                parts.append(s)
            return "-".join(parts)
        s = str(col).strip()
        return "" if s.startswith("Unnamed:") else s

    @staticmethod
    def _clean_excel_dataframe(df: pd.DataFrame) -> pd.DataFrame:
        """清洗 pandas 读取的 DataFrame：
        1. 丢弃整列为空 / 列名全空的列；
        2. 丢弃整行为空的行；
        3. **保留原始类型**（日期仍然是 datetime），不强转字符串，否则日期会被
           Excel 序列号（如 44927）污染。
        """
        if df is None or df.empty:
            return df

        df = df.copy()
        # 只规范化字符串形式的 'nan' / 'NaN'，不 fillna（否则 datetime 列会崩）
        df = df.replace({"nan": None, "NaN": None})

        keep_idx: list[int] = []
        for i in range(df.shape[1]):
            col_name = str(df.columns[i]).strip()
            col_values = df.iloc[:, i]
            is_all_empty = col_values.isna().all() or (
                col_values.astype(str).str.strip() == ""
            ).all()
            if is_all_empty and (not col_name or col_name.startswith("Unnamed:")):
                continue
            if is_all_empty and not col_name:
                continue
            keep_idx.append(i)

        if not keep_idx:
            return df.iloc[0:0]
        df = df.iloc[:, keep_idx]

        df.columns = [
            "" if str(col).startswith("Unnamed:") else str(col).strip()
            for col in df.columns
        ]

        if not df.empty:
            row_is_empty = df.apply(
                lambda row: all(
                    (v is None)
                    or (isinstance(v, float) and pd.isna(v))
                    or (str(v).strip() == "")
                    for v in row
                ),
                axis=1,
            )
            df = df.loc[~row_is_empty]

        return df

    @staticmethod
    def _dataframe_to_records_text(df: pd.DataFrame, *, sheet_name: str) -> str:
        """把 DataFrame 转成面向 LLM 的"记录式"文本：

            工作表：2024
            字段：日期、时间、博总线-正向有功总、博总线-反向有功总、...
            行1：日期=2024-01-01 | 时间=08:00 | 博总线-正向有功总=8608.70 | ...
            行2：日期=2024-01-01 | 时间=16:00 | ...

        - 日期列格式化，不再出现 44927 这种 Excel 序列号；
        - 空值不写成 "key="，降噪；
        - 首行 "字段：..." 会被 _split_csv_document 当表头复用到每个 chunk。
        """
        if df is None or df.empty:
            return ""

        df = df.copy()

        for col in df.columns:
            s = df[col]
            if pd.api.types.is_datetime64_any_dtype(s):
                dt = pd.to_datetime(s, errors="coerce")
                has_time = bool(
                    (
                        (dt.dt.hour.fillna(0) != 0)
                        | (dt.dt.minute.fillna(0) != 0)
                        | (dt.dt.second.fillna(0) != 0)
                    ).any()
                )
                fmt = "%Y-%m-%d %H:%M:%S" if has_time else "%Y-%m-%d"
                df[col] = dt.dt.strftime(fmt)

        cols = [str(c).strip() for c in df.columns]
        non_empty_cols = [c for c in cols if c]

        lines: list[str] = [f"工作表：{sheet_name}"]
        if non_empty_cols:
            lines.append("字段：" + "、".join(non_empty_cols))

        for i, row in enumerate(df.itertuples(index=False, name=None), start=1):
            parts: list[str] = []
            for col, val in zip(cols, row):
                if not col:
                    continue
                if val is None:
                    continue
                if isinstance(val, float) and pd.isna(val):
                    continue
                # 去掉 5566.780000 这种拖尾 0
                if isinstance(val, float):
                    if val.is_integer():
                        val_str = str(int(val))
                    else:
                        val_str = ("%.6f" % val).rstrip("0").rstrip(".")
                else:
                    val_str = str(val).strip()
                if not val_str:
                    continue
                parts.append(f"{col}={val_str}")
            if parts:
                lines.append(f"行{i}：" + " | ".join(parts))

        # 只有工作表/字段两行、没数据行 => 视为空
        if len(lines) <= 2:
            return ""
        return "\n".join(lines)
