import asyncio
from pathlib import Path
from uuid import uuid4

import aiofiles
import markdown
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_core.messages import AIMessage, HumanMessage, SystemMessage
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_text_splitters import Language, RecursiveCharacterTextSplitter
from pypdf import PdfReader
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from core.settings import settings
from db.models.chat import ChatThread
from db.models.knowledge import KnowledgeChunk
from db.session import AsyncSessionLocal
from repository.chat_repo import (
    append_chat_message,
    create_chat_thread,
    get_chat_thread,
    list_recent_unsummarized_messages,
    list_unsummarized_messages,
    mark_messages_summarized,
    update_thread_summary,
)
from repository.knowledge_repo import (
    bulk_insert_chunks,
    clear_chunks_by_file_id,
    create_knowledge_file,
    get_knowledge_file_by_id,
    list_knowledge_files_page,
    soft_delete_knowledge_file,
    update_knowledge_file_status,
)


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".xlsx",
    ".xls",
    ".csv",
    ".md",
    ".markdown",
    ".txt",
}


def get_llm() -> ChatOpenAI:
    return ChatOpenAI(
        model=settings.QWEN_CHAT_MODEL,
        api_key=settings.DASHSCOPE_API_KEY,
        base_url=settings.DASHSCOPE_BASE_URL,
        temperature=0.2,
        max_retries=2,
        timeout=60,
    )


def get_embeddings() -> OpenAIEmbeddings:
    return OpenAIEmbeddings(
        model=settings.QWEN_EMBEDDING_MODEL,
        api_key=settings.DASHSCOPE_API_KEY,
        base_url=settings.DASHSCOPE_BASE_URL,
    )


def estimate_tokens(text: str) -> int:
    if not text:
        return 0
    return max(1, len(text) // 3)


def parse_file(path: str, filename: str) -> list[tuple[str, dict]]:
    file_path = Path(path)
    ext = file_path.suffix.lower()

    if ext == ".pdf":
        reader = PdfReader(str(file_path))
        result = []
        for i, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                result.append((text, {"page": i}))
        return result

    if ext == ".docx":
        doc = DocxDocument(str(file_path))
        blocks = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                blocks.append(text)

        for table_idx, table in enumerate(doc.tables, start=1):
            blocks.append(f"[表格 {table_idx}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))

        text = "\n".join(blocks).strip()
        return [(text, {"section": "document"})] if text else []

    if ext in {".xlsx", ".xls"}:
        xls = pd.ExcelFile(file_path)
        result = []

        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name=sheet_name, dtype=str).fillna("")
            lines = [f"[Sheet] {sheet_name}"]
            headers = [str(c).strip() for c in df.columns.tolist()]
            if headers:
                lines.append(" | ".join(headers))
            for row in df.itertuples(index=False):
                vals = [str(v).strip() for v in row]
                if any(vals):
                    lines.append(" | ".join(vals))
            text = "\n".join(lines).strip()
            if text:
                result.append((text, {"sheet": sheet_name}))
        return result

    if ext == ".csv":
        df = pd.read_csv(file_path, dtype=str).fillna("")
        lines = []
        headers = [str(c).strip() for c in df.columns.tolist()]
        if headers:
            lines.append(" | ".join(headers))
        for row in df.itertuples(index=False):
            vals = [str(v).strip() for v in row]
            if any(vals):
                lines.append(" | ".join(vals))
        text = "\n".join(lines).strip()
        return [(text, {"section": "csv"})] if text else []

    if ext in {".md", ".markdown"}:
        raw = file_path.read_text(encoding="utf-8", errors="ignore")
        html = markdown.markdown(raw)
        text = BeautifulSoup(html, "html.parser").get_text("\n").strip()
        return [(text, {"section": "markdown"})] if text else []

    if ext == ".txt":
        text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
        return [(text, {"section": "txt"})] if text else []

    raise ValueError(f"不支持的文件类型: {ext}")


def build_splitter(filename: str):
    ext = Path(filename).suffix.lower()

    if ext in {".md", ".markdown"}:
        return RecursiveCharacterTextSplitter.from_language(
            language=Language.MARKDOWN,
            chunk_size=settings.KNOWLEDGE_CHUNK_SIZE,
            chunk_overlap=settings.KNOWLEDGE_CHUNK_OVERLAP,
        )

    return RecursiveCharacterTextSplitter(
        chunk_size=settings.KNOWLEDGE_CHUNK_SIZE,
        chunk_overlap=settings.KNOWLEDGE_CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "！", "？", ".", " ", ""],
    )


def split_sections(filename: str, sections: list[tuple[str, dict]]) -> list[Document]:
    splitter = build_splitter(filename)
    docs: list[Document] = []

    for text, meta in sections:
        partial_docs = splitter.create_documents(
            texts=[text],
            metadatas=[meta],
        )
        docs.extend(partial_docs)

    return docs


async def aembed_documents(texts: list[str]) -> list[list[float]]:
    embeddings = get_embeddings()
    vectors: list[list[float]] = []

    batch_size = 10
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        part = await asyncio.to_thread(embeddings.embed_documents, batch)
        vectors.extend(part)

    return vectors


async def aembed_query(text: str) -> list[float]:
    embeddings = get_embeddings()
    return await asyncio.to_thread(embeddings.embed_query, text)


async def create_knowledge_file_and_process(
    db: AsyncSession,
    *,
    folder_id: str | None,
    upload_file,
):
    # 保存上传的文件到本地，并获取保存路径和文件大小
    save_path, size = await save_upload_file(upload_file)

    # 检查文件大小是否超出配置的最大值，超出则抛出异常
    if size > settings.KNOWLEDGE_MAX_FILE_SIZE_MB * 1024 * 1024:
        raise ValueError("文件大小超限")

    # 获取文件扩展名
    ext = Path(upload_file.filename).suffix.lower()
    # 在数据库中创建知识库文件记录
    row = await create_knowledge_file(
        db,
        folder_id=folder_id,
        filename=upload_file.filename,
        ext=ext,
        content_type=upload_file.content_type,
        size=size,
        storage_path=save_path,
    )

    # 异步启动后台解析文件的任务，不阻塞当前请求
    asyncio.create_task(process_knowledge_file(row.id))
    # 返回文件对象
    return row


async def process_knowledge_file(file_id: str):
    # 创建数据库会话
    async with AsyncSessionLocal() as db:
        # 查询文件对象
        file_obj = await get_knowledge_file_by_id(db, file_id)
        if not file_obj:
            return  # 文件不存在，直接返回

        try:
            # 更新文件状态为“解析中”
            await update_knowledge_file_status(
                db,
                file_id=file_id,
                status="parsing",
                error_message=None,
            )

            # 在后台线程中解析文件，获得分段信息
            sections = await asyncio.to_thread(
                parse_file,
                file_obj.storage_path,
                file_obj.filename,
            )

            # 对分段信息进行进一步分割，得到文档片段
            docs = split_sections(file_obj.filename, sections)
            texts = [doc.page_content for doc in docs]  # 文本内容列表
            metas = [doc.metadata for doc in docs]  # 元数据列表

            # 清除该文件已有的切片数据
            await clear_chunks_by_file_id(db, file_id)

            if texts:
                # 生成每个文本的向量嵌入
                vectors = await aembed_documents(texts)

                rows = []
                # 组装要插入的切片对象
                for idx, (text, meta, vector) in enumerate(
                    zip(texts, metas, vectors), start=1
                ):
                    rows.append(
                        KnowledgeChunk(
                            file_id=file_obj.id,
                            folder_id=file_obj.folder_id,
                            chunk_index=idx,
                            content=text,
                            meta_json=meta,
                            embedding=vector,
                        )
                    )

                # 批量插入切片到数据库
                await bulk_insert_chunks(db, rows)

            # 更新文件状态为“可用”，同时更新切片数量
            await update_knowledge_file_status(
                db,
                file_id=file_id,
                status="ready",
                chunk_count=len(texts),
            )

        except Exception as e:
            # 异常时，更新文件状态为“失败”，保存错误信息
            await update_knowledge_file_status(
                db,
                file_id=file_id,
                status="failed",
                error_message=str(e)[:1000],
            )


async def save_upload_file(upload_file) -> tuple[str, int]:
    ext = Path(upload_file.filename or "").suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型: {ext}")

    save_dir = Path(settings.KNOWLEDGE_UPLOAD_DIR)
    save_dir.mkdir(parents=True, exist_ok=True)

    save_name = f"{uuid4().hex}{ext}"
    save_path = save_dir / save_name

    total_size = 0
    async with aiofiles.open(save_path, "wb") as f:
        while chunk := await upload_file.read(1024 * 1024):
            total_size += len(chunk)
            await f.write(chunk)

    await upload_file.close()
    return str(save_path), total_size


async def list_knowledge_files_service(
    db: AsyncSession,
    *,
    page: int,
    page_size: int,
    folder_id: str | None = None,
    keyword: str | None = None,
    status: str | None = None,
):
    rows, pagination = await list_knowledge_files_page(
        db,
        page=page,
        page_size=page_size,
        folder_id=folder_id,
        keyword=keyword,
        status=status,
    )

    return {
        "items": [
            {
                "id": row.id,
                "folder_id": row.folder_id,
                "filename": row.filename,
                "ext": row.ext,
                "status": row.status,
                "size": row.size,
                "chunk_count": row.chunk_count,
                "created_at": row.created_at,
                "updated_at": row.updated_at,
            }
            for row in rows
        ],
        "pagination": pagination,
    }


async def get_knowledge_file_detail_service(
    db: AsyncSession,
    *,
    file_id: str,
):
    row = await get_knowledge_file_by_id(db, file_id)
    if not row:
        return None

    return {
        "id": row.id,
        "folder_id": row.folder_id,
        "filename": row.filename,
        "ext": row.ext,
        "content_type": row.content_type,
        "size": row.size,
        "storage_path": row.storage_path,
        "status": row.status,
        "error_message": row.error_message,
        "chunk_count": row.chunk_count,
        "created_at": row.created_at,
        "updated_at": row.updated_at,
    }


async def delete_knowledge_file_service(
    db: AsyncSession,
    *,
    file_id: str,
):
    row = await get_knowledge_file_by_id(db, file_id)
    if not row:
        return False

    await soft_delete_knowledge_file(db, file_id=file_id)
    return True


def to_langchain_message(role: str, content: str):
    if role == "user":
        return HumanMessage(content=content)
    if role == "assistant":
        return AIMessage(content=content)
    return SystemMessage(content=content)


async def ensure_thread(
    db: AsyncSession, thread_id: str | None, first_question: str
) -> ChatThread:
    if thread_id:
        thread = await get_chat_thread(db, thread_id)
        if thread:
            return thread

    return await create_chat_thread(
        db,
        title=first_question[:30] or "新对话",
    )


async def maybe_summarize_thread(db: AsyncSession, thread_id: str):
    thread = await get_chat_thread(db, thread_id)
    if not thread:
        return

    rows = await list_unsummarized_messages(db, thread_id=thread_id)
    total_tokens = sum(row.tokens_estimate for row in rows) + estimate_tokens(
        thread.summary or ""
    )

    if (
        len(rows) < settings.THREAD_SUMMARY_TRIGGER_MSG_COUNT
        and total_tokens < settings.THREAD_SUMMARY_TRIGGER_TOKEN_EST
    ):
        return

    keep_n = settings.THREAD_KEEP_RECENT_MSGS
    if len(rows) <= keep_n:
        return

    to_summarize = rows[:-keep_n]
    if not to_summarize:
        return

    transcript = "\n".join(f"[{row.role}] {row.content}" for row in to_summarize)

    llm = get_llm()
    messages = [
        SystemMessage(
            content=(
                "你负责压缩对话历史。"
                "请保留：用户目标、约束条件、已确认结论、未解决问题、关键上下文。"
                "输出简洁中文摘要。"
            )
        ),
        HumanMessage(
            content=(
                f"现有摘要：\n{thread.summary or '无'}\n\n"
                f"新增待压缩对话：\n{transcript}"
            )
        ),
    ]
    resp = await llm.ainvoke(messages)
    summary = resp.content if isinstance(resp.content, str) else str(resp.content)

    await update_thread_summary(db, thread_id=thread_id, summary=summary)
    await mark_messages_summarized(db, message_ids=[row.id for row in to_summarize])


async def load_thread_context(db: AsyncSession, thread_id: str):
    thread = await get_chat_thread(db, thread_id)
    if not thread:
        return None, []

    rows = await list_recent_unsummarized_messages(
        db,
        thread_id=thread_id,
        limit=settings.THREAD_KEEP_RECENT_MSGS,
    )
    messages = [to_langchain_message(row.role, row.content) for row in rows]
    return thread.summary, messages


async def retrieve_chunks(
    db: AsyncSession,
    *,
    question: str,
    folder_ids: list[str] | None = None,
    file_ids: list[str] | None = None,
    top_k: int | None = None,
):
    top_k = top_k or settings.KNOWLEDGE_RETRIEVE_TOP_K
    query_vector = await aembed_query(question)
    distance = KnowledgeChunk.embedding.cosine_distance(query_vector)

    stmt = (
        select(
            KnowledgeChunk.id.label("chunk_id"),
            KnowledgeChunk.file_id,
            KnowledgeChunk.folder_id,
            KnowledgeChunk.chunk_index,
            KnowledgeChunk.content,
            KnowledgeChunk.meta_json,
            distance.label("score"),
        )
        .where(
            KnowledgeChunk.is_deleted == 0,
        )
        .order_by(distance.asc())
        .limit(top_k)
    )

    if folder_ids:
        stmt = stmt.where(KnowledgeChunk.folder_id.in_(folder_ids))
    if file_ids:
        stmt = stmt.where(KnowledgeChunk.file_id.in_(file_ids))

    rows = (await db.execute(stmt)).mappings().all()
    if not rows:
        return []

    file_ids_all = list({row["file_id"] for row in rows})
    file_stmt = select(
        KnowledgeChunk.file_id,
    )
    _ = file_stmt  # 占位，保持风格一致

    # 这里直接用额外查询文件名
    from db.models.knowledge import KnowledgeFile

    file_map_stmt = select(KnowledgeFile.id, KnowledgeFile.filename).where(
        KnowledgeFile.id.in_(file_ids_all),
        KnowledgeFile.is_deleted == 0,
    )
    file_rows = (await db.execute(file_map_stmt)).all()
    file_map = {row[0]: row[1] for row in file_rows}

    result = []
    for row in rows:
        result.append(
            {
                "chunk_id": row["chunk_id"],
                "file_id": row["file_id"],
                "folder_id": row["folder_id"],
                "chunk_index": row["chunk_index"],
                "content": row["content"],
                "meta_json": row["meta_json"] or {},
                "filename": file_map.get(row["file_id"], "未知文件"),
                "score": float(row["score"]),
            }
        )
    return result


def build_raw_context(rows) -> str:
    blocks = []
    for idx, row in enumerate(rows, start=1):
        meta = row["meta_json"] or {}
        loc_parts = []
        if meta.get("page"):
            loc_parts.append(f"page={meta['page']}")
        if meta.get("sheet"):
            loc_parts.append(f"sheet={meta['sheet']}")
        if meta.get("section"):
            loc_parts.append(f"section={meta['section']}")

        loc_text = ", ".join(loc_parts) if loc_parts else "unknown"

        blocks.append(
            f"[证据{idx}]\n"
            f"文件: {row['filename']}\n"
            f"位置: {loc_text}\n"
            f"内容:\n{row['content']}"
        )
    return "\n\n".join(blocks)


async def compress_context_if_needed(question: str, rows) -> str:
    raw_context = build_raw_context(rows)
    if len(raw_context) <= settings.MAX_CONTEXT_CHARS:
        return raw_context

    groups = []
    current_group = []
    current_chars = 0

    for row in rows:
        text = (
            f"文件: {row['filename']}\n"
            f"元信息: {row['meta_json']}\n"
            f"内容:\n{row['content']}\n"
        )
        if current_chars + len(text) > settings.CONTEXT_GROUP_CHARS and current_group:
            groups.append(current_group)
            current_group = []
            current_chars = 0

        current_group.append(text)
        current_chars += len(text)

    if current_group:
        groups.append(current_group)

    llm = get_llm()
    partial_summaries = []

    for idx, group in enumerate(groups, start=1):
        group_text = "\n\n".join(group)
        resp = await llm.ainvoke(
            [
                SystemMessage(
                    content=(
                        "你负责压缩知识证据。"
                        "只保留与用户问题直接相关的事实、步骤、参数、结论。"
                        "不要扩写，不要推断。"
                    )
                ),
                HumanMessage(
                    content=f"用户问题:\n{question}\n\n第{idx}组证据:\n{group_text}"
                ),
            ]
        )
        partial_summaries.append(
            resp.content if isinstance(resp.content, str) else str(resp.content)
        )

    final_resp = await llm.ainvoke(
        [
            SystemMessage(
                content="你负责合并多组知识证据摘要，只输出最终可用于回答问题的精炼证据。"
            ),
            HumanMessage(
                content=f"用户问题:\n{question}\n\n多组摘要:\n"
                + "\n\n".join(partial_summaries)
            ),
        ]
    )
    return (
        final_resp.content
        if isinstance(final_resp.content, str)
        else str(final_resp.content)
    )


async def ask_with_rag_service(
    db: AsyncSession,
    *,
    question: str,
    thread_id: str | None = None,
    folder_ids: list[str] | None = None,
    file_ids: list[str] | None = None,
    top_k: int = 5,
):
    thread = await ensure_thread(db, thread_id, question)

    user_msg = await append_chat_message(
        db,
        thread_id=thread.id,
        role="user",
        content=question,
        tokens_estimate=estimate_tokens(question),
    )

    await maybe_summarize_thread(db, thread.id)
    thread_summary, history_messages = await load_thread_context(db, thread.id)

    retrieved_rows = await retrieve_chunks(
        db,
        question=question,
        folder_ids=folder_ids,
        file_ids=file_ids,
        top_k=top_k,
    )

    context_text = (
        await compress_context_if_needed(question, retrieved_rows)
        if retrieved_rows
        else "无可用证据"
    )

    prompt = ChatPromptTemplate.from_messages(
        [
            ("system", "你是运行智脑 AI 助手。"),
            (
                "system",
                "必须优先依据知识库证据回答；如果知识库证据不足，明确说“当前知识库中未找到足够依据”。不要编造。",
            ),
            ("system", "线程历史摘要：\n{thread_summary}"),
            MessagesPlaceholder("history"),
            ("human", "知识库证据：\n{context}\n\n当前问题：\n{question}"),
        ]
    )

    llm = get_llm()
    chain = prompt | llm
    ai_resp = await chain.ainvoke(
        {
            "thread_summary": thread_summary or "无",
            "history": history_messages,
            "context": context_text,
            "question": question,
        }
    )

    answer = (
        ai_resp.content if isinstance(ai_resp.content, str) else str(ai_resp.content)
    )

    references = [
        {
            "chunk_id": row["chunk_id"],
            "file_id": row["file_id"],
            "filename": row["filename"],
            "chunk_index": row["chunk_index"],
            "meta": row["meta_json"],
            "score": row["score"],
        }
        for row in retrieved_rows
    ]

    await append_chat_message(
        db,
        thread_id=thread.id,
        role="assistant",
        content=answer,
        parent_message_id=user_msg.id,
        tokens_estimate=estimate_tokens(answer),
        meta_json={"references": references},
    )

    return {
        "thread_id": thread.id,
        "answer": answer,
        "references": references,
    }
